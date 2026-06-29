"""Genera el documento Word maestro del sistema de capacitación e inspecciones post-sismo."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "diseno" / "DOCUMENTO-MAESTRO-SISTEMA.docx"


def _h(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _p(doc: Document, text: str, bold: bool = False) -> None:
    run = doc.add_paragraph().add_run(text)
    run.bold = bold


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


def _screen_block(
    doc: Document,
    *,
    titulo: str,
    ruta: str,
    roles: str,
    proposito: str,
    contenido: list[str],
    botones: list[str],
    validaciones: list[str] | None = None,
    salidas: list[str] | None = None,
) -> None:
    _h(doc, titulo, 3)
    _p(doc, f"Ruta URL: {ruta}")
    _p(doc, f"Roles: {roles}")
    _p(doc, f"Propósito: {proposito}", bold=True)
    _p(doc, "Contenido en pantalla (muestra):")
    _bullets(doc, contenido)
    _p(doc, "Botones y acciones:")
    _bullets(doc, botones)
    if validaciones:
        _p(doc, "Validaciones:")
        _bullets(doc, validaciones)
    if salidas:
        _p(doc, "Salidas / navegación:")
        _bullets(doc, salidas)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # Portada
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("DOCUMENTO MAESTRO DEL PROYECTO\n")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    t.add_run(
        "Sistema de Capacitación e Inspecciones\n"
        "Evaluación Rápida de Daños en Edificaciones (ERD)\n"
        "Post-terremoto · Caracas\n\n"
    ).font.size = Pt(14)
    t.add_run(f"Versión 1.1 · {date.today().isoformat()}\n").font.size = Pt(11)
    t.add_run("Marco: Django · PostgreSQL · Metodología ANIH Planilla V.8").font.size = Pt(11)
    doc.add_page_break()

    _h(doc, "1. Resumen ejecutivo", 1)
    _p(
        doc,
        "El proyecto consiste en un sistema web institucional para: (a) registrar y administrar "
        "participantes de cursos de inducción/capacitación de inspectores; (b) habilitar inspectores "
        "certificados; (c) capturar en campo la Evaluación Rápida de Daños en Edificaciones según "
        "la Planilla y Manual de Campo V.8 y el Boletín 61 ANIH; (d) calcular automáticamente "
        "riesgos y etiquetas de acceso (verde, amarilla, roja); (e) generar PDF de la planilla "
        "y reportes para coordinación.",
    )

    _h(doc, "2. Objetivo y alcance", 1)
    _h(doc, "2.1 Objetivo general", 2)
    _p(doc, "Sistematizar la información de cursos e inspecciones post-evento sísmico con trazabilidad, "
         "seguridad de datos personales y salidas oficiales (PDF, listados, mapas).")
    _h(doc, "2.2 Alcance incluido", 2)
    _bullets(doc, [
        "Catálogo de sedes (universidades, auditorios) — alta única.",
        "Calendario de sesiones de inducción con carga en lote (no curso por curso).",
        "Inscripción pública: elegir sede → fecha/hora → datos personales (10 variables).",
        "QR por sede para inscripción en el lugar.",
        "Control de asistencia y emisión de certificado/habilitación.",
        "Carga digital de inspecciones ERD (pasos 0 a 7 de la planilla V.8).",
        "Reglas automáticas de corte (riesgo externo/estructural C → etiqueta roja).",
        "Adjuntos: fotos, croquis, observaciones.",
        "Generación de PDF de planilla y etiqueta de acceso.",
        "Panel de administración y reportes.",
        "Gestión de usuarios, roles y permisos.",
    ])
    _h(doc, "2.3 Fuera de alcance inicial (fases posteriores)", 2)
    _bullets(doc, [
        "App móvil nativa (fase 1 usa web responsive).",
        "Inspección detallada posterior (solo se registra como acción recomendada).",
        "Integración GIS municipal en tiempo real.",
        "Firma electrónica avanzada con certificado digital.",
    ])

    _h(doc, "2.4 Operación masiva — simplificación de capacitación", 2)
    _p(
        doc,
        "Las inducciones se realizan de forma masiva en universidades y auditorios, a menudo con el "
        "mismo equipo coordinador. El sistema NO exige crear manualmente un «curso» por cada sesión. "
        "En su lugar usa tres conceptos:",
    )
    _table(doc, ["Concepto", "Qué es", "Frecuencia de alta"], [
        ["Campaña", "Programa general (ej. «Capacitación post-sismo Caracas 2026»)", "Una vez"],
        ["Sede", "Universidad, auditorio o institución (lugar fijo)", "Una vez por sede; catálogo reutilizable"],
        ["Sesión", "Sede + fecha + hora (equivale al «curso» del formulario)", "En lote o import Excel"],
    ])
    _p(doc, "En pantalla el usuario puede seguir viendo la palabra «curso» o «inducción», pero "
         "técnicamente es una sesión en una sede.", bold=True)
    _h(doc, "2.4.1 Formas de crear sesiones (coordinador)", 3)
    _table(doc, ["Método", "Uso", "Ejemplo"], [
        ["Carga en lote", "Operación masiva principal", "UCV, 1–15 jul, Lun–Vie, 09:00 → 11 sesiones"],
        ["Importar Excel", "Planificación semanal", "Columnas: sede, fecha, hora, cupo"],
        ["Duplicar semana", "Repetir patrón", "Copiar semana pasada en misma sede"],
        ["Sesión individual", "Excepciones", "Una fecha extra en auditorio X"],
        ["QR por sede", "Inscripción in situ", "Participante elige fecha disponible en esa sede"],
    ])
    _h(doc, "2.4.2 Responsable por defecto", 3)
    _p(doc, "El instructor/coordinador se asigna a nivel de Campaña o Sede (misma persona en todas "
         "las sesiones de UCV), no se repite en cada sesión salvo excepción.")

    _h(doc, "3. Marco de trabajo tecnológico", 1)
    _table(doc, ["Capa", "Tecnología", "Justificación"], [
        ["Backend", "Django 5.x (Python)", "Robusto, Admin integrado, auth, ORM, escalable"],
        ["Base de datos", "PostgreSQL (prod) / SQLite (dev)", "Una sola BD para cursos e inspecciones"],
        ["Frontend", "Django Templates + Bootstrap 5", "Formularios, wizard, responsive móvil"],
        ["Formularios", "django-crispy-forms", "UX consistente en pasos de inspección"],
        ["Wizard inspección", "django-formtools o vistas por paso", "Flujo paso a paso con reglas"],
        ["PDF", "WeasyPrint (HTML→PDF)", "Planilla V.8, fotos, etiquetas"],
        ["Correos", "SMTP (Brevo / institucional)", "Confirmaciones, invitaciones, reset clave"],
        ["Archivos", "Django FileField + media/", "Fotos y croquis"],
        ["Admin", "Django Admin + vistas custom", "Coordinadores y consulta"],
        ["Control versiones", "Git + GitHub", "Despliegue continuo"],
        ["Hosting gratuito", "Render / Railway / VPS Oracle Free", "Demo y piloto"],
        ["Seguridad", "HTTPS, CSRF, permisos por rol", "Datos personales C.I., correo, teléfono"],
    ])

    _h(doc, "4. Marco metodológico de inspección", 1)
    _p(doc, "El sistema digitaliza el instrumento oficial:")
    _bullets(doc, [
        "Instrumento de Inspección — Planilla y Manual de Campo V.8.",
        "Evaluación de Daños — Boletín 61 ANIH.",
        "Metodología: Evaluación Rápida de Daños (ERD) en edificaciones.",
        "Resultado: etiqueta Verde (acceso permitido), Amarilla (restringido), Roja (no permitido).",
        "Solo inspectores con certificado de capacitación pueden cargar inspecciones.",
    ])
    _h(doc, "4.1 Pasos de la planilla digitalizados", 2)
    _table(doc, ["Paso", "Sección", "Regla de corte"], [
        ["0", "Encabezado (planilla Nº, evento, fechas, horas)", "—"],
        ["1", "Información general y edificación", "—"],
        ["2", "Inspección externa → Riesgo externo A/B/C", "Si C → ir a paso 6, etiqueta roja"],
        ["3", "Piso crítico — daño severo/completo (N)", "Si N≥1 → ir a paso 6, etiqueta roja"],
        ["4", "Daño moderado en piso crítico (%)", "Continúa si no hay severo"],
        ["5", "Componentes no estructurales", "Riesgo componentes A/B/C"],
        ["6", "Riesgo final = máximo(2,3,4,5) + etiqueta", "Verde / Amarilla / Roja"],
        ["7", "Acciones recomendadas", "Checklist"],
        ["Anexo", "Observaciones, croquis, fotos, firmas", "—"],
    ])

    _h(doc, "5. Roles de usuario y permisos", 1)
    _table(doc, ["Rol", "Usuario propio", "Permisos principales"], [
        ["Administrador", "Sí (1-2 personas)", "Usuarios, configuración, respaldos, todo"],
        ["Coordinador", "Sí", "Sedes, calendario, sesiones en lote, asistencia, inspecciones, PDF, reportes"],
        ["Registrador", "Sí (opcional)", "Cargar sesiones (lote/Excel), inscripciones y asistencia del día"],
        ["Inspector certificado", "Sí (cada inspector)", "Crear/editar sus inspecciones, subir fotos, PDF propio"],
        ["Consulta / Gerencia", "Sí", "Solo lectura, exportar reportes"],
        ["Público (sin login)", "No", "Inscripción: elegir sede y sesión (o QR por sede)"],
    ])
    _p(doc, "Recomendación: NO centralizar un solo usuario para todos los inspectores. "
         "Cada inspector debe tener credenciales propias para trazabilidad legal y operativa.")

    _h(doc, "6. Arquitectura del sistema", 1)
    _p(doc, "Aplicación monolítica Django con módulos (apps) separados y una sola base de datos PostgreSQL:")
    _bullets(doc, [
        "apps/capacitacion — Campaña, Sede, Sesión, inscripciones, asistencia, certificados.",
        "apps/participantes — Personas (C.I., contacto, procedencia).",
        "apps/inspecciones — Planilla ERD completa, adjuntos, PDF.",
        "apps/cuentas — Perfiles de usuario vinculados a participante.",
        "apps/reportes — Dashboards, export Excel, mapa.",
    ])
    _p(doc, "Flujo de datos: Campaña → Sede → Sesión → Inscripción → Participante (certificado) "
         "→ User inspector → Inspección → PDF.")

    _h(doc, "7. Modelo de datos (resumen)", 1)
    _table(doc, ["Entidad", "Campos clave", "Relación"], [
        ["Campaña", "nombre, responsable_default, activa", "1:N Sedes (opcional agrupación)"],
        ["Sede", "nombre, tipo, dirección, municipio, QR slug", "1:N Sesiones"],
        ["Sesión", "fecha, hora, cupo, estado, responsable (hereda sede)", "N:1 Sede; 1:N Inscripciones"],
        ["Participante", "cédula, apellidos, nombres, teléfono, correo", "1:N Inscripciones; 1:1 User (opcional)"],
        ["Inscripción", "profesión, procedencia, asistió, certificado", "N:1 Sesión, N:1 Participante"],
        ["Inspección", "planilla_n, evento, fechas, estado borrador/cerrada", "N:1 Inspector(User)"],
        ["Edificación", "ubicación, coordenadas, uso, pisos, material", "1:1 Inspección"],
        ["InspecciónExterna", "5 aspectos, riesgo_externo", "1:1 Inspección"],
        ["PisoCrítico / Elementos", "conteos, % moderado, N severo", "1:N Inspección"],
        ["ComponentesNoEstruct.", "losas, paredes, instalaciones…", "1:N Inspección"],
        ["ResultadoAcceso", "riesgo_final, etiqueta color", "1:1 Inspección"],
        ["AcciónRecomendada", "apuntalar, acordonar, etc.", "N:M Inspección"],
        ["Adjunto", "foto, croquis", "N:1 Inspección"],
    ])

    _h(doc, "8. Variables del formulario de inscripción", 1)
    _p(doc, "Las variables 1–3 no se teclean: el participante elige Sede y Sesión (fecha/hora) "
         "en los pasos previos del formulario o vienen precargadas por QR.")
    _table(doc, ["#", "Variable", "Tipo en sistema", "Validación"], [
        ["1", "Lugar de Inducción", "Select sede (catálogo)", "Automático al elegir sede"],
        ["2", "Fecha de Inducción", "Select sesión disponible", "Solo fechas futuras con cupo"],
        ["3", "Hora", "Parte de la sesión elegida", "Automático"],
        ["4", "C.I.", "Texto (V/E + dígitos)", "Único por sesión, formato"],
        ["5", "Apellidos", "Texto", "Obligatorio"],
        ["6", "Nombres", "Texto", "Obligatorio"],
        ["7", "Teléfono Celular", "Texto", "Obligatorio"],
        ["8", "Correo electrónico", "Email", "Formato email"],
        ["9", "Profesión/Ocupación", "Select + Otro", "Si Otro → texto"],
        ["10", "Procedencia", "Texto", "Universidad/institución/empresa"],
    ])

    _h(doc, "9. Especificación de pantallas", 1)
    _p(doc, "A continuación se detalla cada pantalla con contenido de ejemplo, botones y navegación.")

    # --- PANTALLAS PÚBLICAS ---
    _h(doc, "9.1 Pantallas públicas (sin login)", 2)

    _screen_block(
        doc,
        titulo="P-01 · Inicio público",
        ruta="/",
        roles="Público",
        proposito="Punto de entrada institucional.",
        contenido=[
            "Logo institución / programa post-sismo.",
            "Texto: «Sistema de Capacitación e Inspección — Evaluación Rápida de Daños».",
            "Enlaces: Inscribirse a inducción | Iniciar sesión (inspectores/coordinadores).",
            "Aviso de privacidad de datos personales.",
        ],
        botones=["Inscribirse", "Iniciar sesión"],
        salidas=["→ P-02 Inscripción (elegir sede)", "→ P-02b QR sede", "→ L-01 Login"],
    )

    _screen_block(
        doc,
        titulo="P-02 · Inscripción — Paso 1: elegir sede y sesión",
        ruta="/inscripcion/",
        roles="Público",
        proposito="El participante elige dónde y cuándo sin que el coordinador cree enlaces uno a uno.",
        contenido=[
            "Paso 1 — Sede: lista desplegable (Universidad Central, UNEXPO, Auditorio Municipal…).",
            "Paso 2 — Sesión: calendario o lista de fechas/horas disponibles en esa sede con cupo.",
            "Ejemplo: «UCV — 15/07/2026 09:00 (quedan 42 cupos)».",
            "Alternativa: llegada por QR de sede → sede ya seleccionada, solo elige fecha/hora.",
        ],
        botones=["Continuar → datos personales", "Cambiar sede"],
        validaciones=["Sesión con cupo disponible", "Sesión no cancelada"],
        salidas=["→ P-02c Formulario datos"],
    )

    _screen_block(
        doc,
        titulo="P-02b · Inscripción por QR de sede",
        ruta="/inscripcion/sede/<slug>/",
        roles="Público",
        proposito="Cartel QR en la universidad/auditorio; evita buscar sede manualmente.",
        contenido=[
            "Encabezado fijo: «Inducción — Universidad Central de Venezuela».",
            "Solo muestra sesiones de esa sede (hoy y próximas fechas).",
            "Ideal para inscripción en el pasillo el día del evento.",
        ],
        botones=["Elegir sesión", "Continuar"],
        salidas=["→ P-02c Formulario datos"],
    )

    _screen_block(
        doc,
        titulo="P-02c · Inscripción — Paso 2: datos personales",
        ruta="/inscripcion/sesion/<id>/",
        roles="Público",
        proposito="Capturar variables 4–10; variables 1–3 ya definidas por la sesión elegida.",
        contenido=[
            "Resumen fijo: Sede, fecha y hora (no editables).",
            "Formulario: C.I., Apellidos, Nombres, Teléfono, Correo, Profesión (select), Procedencia.",
            "Checkbox: «Acepto el uso de mis datos para fines de capacitación e inspección».",
        ],
        botones=["Enviar inscripción", "← Cambiar sesión"],
        validaciones=[
            "C.I. no duplicada en la misma sesión.",
            "Correo y teléfono con formato válido.",
            "Si Profesión = Otro, campo texto obligatorio.",
        ],
        salidas=[
            "→ P-03 Confirmación.",
            "Correo opcional de confirmación.",
        ],
    )

    _screen_block(
        doc,
        titulo="P-03 · Confirmación de inscripción",
        ruta="/inscripcion/confirmacion/",
        roles="Público",
        proposito="Confirmar registro exitoso.",
        contenido=[
            "«Su inscripción fue registrada correctamente.»",
            "Resumen: nombre, sede, fecha, hora.",
            "Código de inscripción (ej. INS-2026-1847).",
            "Indicaciones: presentar C.I. el día de la inducción.",
        ],
        botones=["Volver al inicio", "Añadir a calendario (opcional)"],
        salidas=["→ P-01"],
    )

    # --- LOGIN ---
    _h(doc, "9.2 Autenticación", 2)

    _screen_block(
        doc,
        titulo="L-01 · Iniciar sesión",
        ruta="/cuentas/login/",
        roles="Todos los usuarios internos",
        proposito="Acceso seguro al sistema.",
        contenido=[
            "Campos: Usuario (o correo) y Contraseña.",
            "Enlace: ¿Olvidó su contraseña?",
        ],
        botones=["Ingresar"],
        validaciones=["Credenciales válidas", "Usuario activo"],
        salidas=[
            "→ D-01 Dashboard según rol.",
            "→ L-02 Recuperar contraseña.",
        ],
    )

    _screen_block(
        doc,
        titulo="L-02 · Recuperar contraseña",
        ruta="/cuentas/password-reset/",
        roles="Usuarios registrados",
        proposito="Enviar enlace por correo para restablecer clave.",
        contenido=["Campo: Correo electrónico registrado."],
        botones=["Enviar enlace", "Volver a login"],
        salidas=["Correo con token → formulario nueva contraseña."],
    )

    # --- COORDINADOR ---
    _h(doc, "9.3 Panel coordinador / administrador", 2)

    _screen_block(
        doc,
        titulo="D-01 · Dashboard coordinador",
        ruta="/panel/",
        roles="Coordinador, Administrador",
        proposito="Vista general de operación.",
        contenido=[
            "KPIs: Sesiones esta semana | Inscritos | Certificados | Inspecciones del mes.",
            "Gráfico: inscripciones por sede y por profesión.",
            "Calendario visual: sesiones próximas (vista mes/semana).",
            "Tabla: sesiones de hoy (asistencia rápida).",
            "Tabla: últimas inspecciones (planilla Nº, dirección, etiqueta).",
        ],
        botones=[
            "Calendario de inducciones",
            "Cargar sesiones en lote",
            "Gestionar sedes",
            "Sesiones de hoy (asistencia)",
            "Ver participantes",
            "Ver inspecciones",
            "Exportar reportes",
            "Administrar usuarios (solo Admin)",
        ],
        salidas=["Navegación a CAL-*, SED-*, SES-*, PAR-*, INS-*"],
    )

    _screen_block(
        doc,
        titulo="SED-01 · Catálogo de sedes",
        ruta="/sedes/",
        roles="Coordinador, Registrador",
        proposito="Alta única de universidades y auditorios; no repetir el nombre cada vez.",
        contenido=[
            "Tabla: Nombre | Tipo (Universidad/Auditorio/Otro) | Municipio | Sesiones activas | QR.",
            "Ejemplo: «Universidad Central de Venezuela | Universidad | Libertador | 12 | [QR]».",
        ],
        botones=["+ Nueva sede", "Editar", "Descargar QR sede", "Ver sesiones de esta sede"],
        salidas=["→ SED-02 Alta/edición sede", "→ CAL-01 Calendario filtrado por sede"],
    )

    _screen_block(
        doc,
        titulo="SED-02 · Alta / edición de sede",
        ruta="/sedes/nueva/ | /sedes/<id>/editar/",
        roles="Coordinador",
        proposito="Registrar lugar de inducción reutilizable.",
        contenido=[
            "Nombre, tipo, dirección, municipio, ciudad, estado.",
            "Responsable por defecto (mismo instructor en todas las sesiones de esta sede).",
            "Slug para URL/QR (ej. ucv).",
            "Observaciones (acceso, piso, contacto en sede).",
        ],
        botones=["Guardar", "Cancelar", "Generar QR"],
        salidas=["→ SED-01"],
    )

    _screen_block(
        doc,
        titulo="CAL-01 · Calendario de inducciones (sesiones)",
        ruta="/capacitacion/calendario/",
        roles="Coordinador, Registrador",
        proposito="Vista central de todas las sesiones; sustituye el listado de «cursos» uno a uno.",
        contenido=[
            "Vista mes/semana/lista con color por sede.",
            "Cada bloque: Sede | Fecha | Hora | Inscritos/Cupo | Estado.",
            "Filtros: sede, campaña, rango de fechas, estado.",
            "Ejemplo: «12/07 UCV 09:00 — 38/80 inscritos — Programada».",
        ],
        botones=[
            "+ Cargar sesiones en lote",
            "+ Sesión individual (excepción)",
            "Importar Excel",
            "Duplicar semana",
            "Exportar calendario",
            "Clic en sesión → detalle",
        ],
        salidas=["→ CAL-02 Carga en lote", "→ SES-03 Detalle sesión"],
    )

    _screen_block(
        doc,
        titulo="CAL-02 · Cargar sesiones en lote",
        ruta="/capacitacion/sesiones/lote/",
        roles="Coordinador, Registrador",
        proposito="Crear decenas de sesiones en una sola operación (caso masivo universidades).",
        contenido=[
            "Sede: [Universidad Central ▼]",
            "Campaña: [Post-sismo Caracas 2026 ▼] (opcional)",
            "Desde fecha / Hasta fecha.",
            "Días de la semana: ☑ Lun ☑ Mar ☑ Mié ☑ Jue ☑ Vie ☐ Sáb ☐ Dom.",
            "Horarios: 09:00 (botón + agregar otro turno ej. 14:00).",
            "Cupo por sesión (opcional). Responsable: hereda de sede (editable).",
            "Vista previa: «Se crearán 22 sesiones» antes de confirmar.",
        ],
        botones=["Vista previa", "Crear sesiones", "Cancelar"],
        validaciones=["Al menos un día y un horario", "Rango de fechas válido"],
        salidas=["→ CAL-01 Calendario con sesiones nuevas"],
    )

    _screen_block(
        doc,
        titulo="CAL-03 · Importar sesiones (Excel)",
        ruta="/capacitacion/sesiones/importar/",
        roles="Coordinador",
        proposito="Planificación masiva desde hoja de cálculo.",
        contenido=[
            "Plantilla descargable: columnas sede_slug, fecha, hora, cupo, observaciones.",
            "Arrastrar archivo .xlsx → validación → resumen de filas OK / con error.",
        ],
        botones=["Descargar plantilla", "Subir archivo", "Importar válidas"],
        salidas=["→ CAL-01"],
    )

    _screen_block(
        doc,
        titulo="SES-03 · Detalle de sesión (antes «detalle de curso»)",
        ruta="/capacitacion/sesiones/<id>/",
        roles="Coordinador, Registrador",
        proposito="Lista de inscritos del día; asistencia y certificación.",
        contenido=[
            "Encabezado: Sede, fecha, hora, responsable, enlace/QR de inscripción.",
            "Tabla inscritos: C.I. | Nombre | Profesión | Procedencia | Asistió | Certificado.",
            "Contadores: 45 inscritos, 38 asistieron, 35 certificados.",
            "Modo «día del evento»: botones grandes para marcar asistencia rápida.",
        ],
        botones=[
            "+ Inscripción manual (walk-in)",
            "Importar inscritos Excel",
            "Marcar todos asistieron",
            "Certificar seleccionados",
            "Exportar lista / PDF lista",
            "Imprimir QR sede",
            "Crear usuarios inspectores (lote)",
        ],
        salidas=["→ PAR-02 Ficha participante"],
    )

    _screen_block(
        doc,
        titulo="SES-04 · Sesiones de hoy (asistencia rápida)",
        ruta="/capacitacion/hoy/",
        roles="Coordinador, Registrador",
        proposito="Pantalla operativa el día de capacitación masiva.",
        contenido=[
            "Solo sesiones con fecha = hoy, agrupadas por sede.",
            "Tarjetas grandes: UCV 09:00 — [Abrir lista] | UNEXPO 14:00 — [Abrir lista].",
            "Contador en vivo de asistencias.",
        ],
        botones=["Abrir lista", "Buscar por C.I. (check-in rápido)"],
        salidas=["→ SES-03 Detalle sesión"],
    )

    _screen_block(
        doc,
        titulo="PAR-01 · Listado de participantes",
        ruta="/participantes/",
        roles="Coordinador",
        proposito="Base maestra de personas.",
        contenido=[
            "Búsqueda por C.I., apellido, correo.",
            "Tabla: C.I. | Nombre completo | Profesión | Sesiones / sedes | Certificado | Usuario sistema.",
        ],
        botones=["Ver ficha", "Exportar Excel"],
        salidas=["→ PAR-02 Ficha participante"],
    )

    _screen_block(
        doc,
        titulo="PAR-02 · Ficha participante",
        ruta="/participantes/<id>/",
        roles="Coordinador",
        proposito="Historial del participante.",
        contenido=[
            "Datos personales.",
            "Historial de sesiones e inscripciones por sede.",
            "Estado: Certificado sí/no, Usuario vinculado sí/no.",
            "Lista de inspecciones realizadas (si es inspector).",
        ],
        botones=[
            "Editar datos",
            "Crear usuario inspector",
            "Enviar invitación por correo",
            "Ver inspecciones",
        ],
        salidas=["→ INS-01 Listado inspecciones filtrado"],
    )

    # --- INSPECTOR ---
    _h(doc, "9.4 Módulo inspector — inspecciones ERD (wizard)", 2)

    _screen_block(
        doc,
        titulo="INS-01 · Mis inspecciones (inspector)",
        ruta="/inspecciones/mis/",
        roles="Inspector certificado",
        proposito="Listar inspecciones propias.",
        contenido=[
            "Tabla: Planilla Nº | Fecha | Dirección | Etiqueta | Estado (borrador/cerrada).",
            "Filtro por fecha y etiqueta.",
        ],
        botones=["+ Nueva inspección", "Continuar borrador", "Ver / PDF"],
        validaciones=["Solo si participante tiene certificado=True"],
        salidas=["→ INS-W0 Wizard"],
    )

    _screen_block(
        doc,
        titulo="INS-W0 · Paso 0 — Encabezado",
        ruta="/inspecciones/nueva/paso-0/",
        roles="Inspector",
        proposito="Datos iniciales de la planilla.",
        contenido=[
            "Planilla Nº: auto-generado (ej. ERD-2026-0042).",
            "Tipo de evento: Sismo / Deslizamiento / Asentamiento / Explosión / Otro.",
            "Fecha del evento: 04/11/2019 (ejemplo).",
            "Fecha inspección: hoy. Hora inicio / Hora fin.",
            "Inspector 1: usuario actual (auto). Inspector 2: select otro certificado.",
        ],
        botones=["Guardar borrador", "Siguiente →"],
        salidas=["→ INS-W1"],
    )

    _screen_block(
        doc,
        titulo="INS-W1 · Paso 1 — Información general y edificación",
        ruta="/inspecciones/<id>/paso-1/",
        roles="Inspector",
        proposito="Ubicación y datos del edificio.",
        contenido=[
            "Ubicación: Estado, Ciudad, Municipio, Sector/Calle.",
            "Coordenadas: Latitud, Longitud (botón «Usar mi ubicación»).",
            "Edificación: Uso predominante (checkbox único), Nº personas, Nº pisos, "
            "semisótanos, sótanos, año construcción.",
            "Material estructura: Concreto / Acero / Mampostería formal / informal / Otro.",
        ],
        botones=["← Anterior", "Guardar borrador", "Siguiente →"],
        salidas=["→ INS-W2"],
    )

    _screen_block(
        doc,
        titulo="INS-W2 · Paso 2 — Inspección externa",
        ruta="/inspecciones/<id>/paso-2/",
        roles="Inspector",
        proposito="Evaluar riesgo externo sin ingresar.",
        contenido=[
            "Tabla 5 filas × 3 columnas (Bajo/Medio/Alto) según planilla:",
            "Colapso estructura | Peligro aledaños | Peligro geológico | Asentamiento | Inclinación.",
            "Resultado calculado: Riesgo Externo A / B / C (solo lectura).",
            "ALERTA si C: «No continúe inspección interna. Vaya al paso 6 — Etiqueta ROJA.»",
        ],
        botones=["← Anterior", "Guardar borrador", "Siguiente →", "Ir a resultado (si C)"],
        validaciones=["Debe seleccionarse una opción por fila"],
        salidas=["→ INS-W3 si A o B", "→ INS-W6 si C"],
    )

    _screen_block(
        doc,
        titulo="INS-W3 · Paso 3 — Piso crítico / daño severo",
        ruta="/inspecciones/<id>/paso-3/",
        roles="Inspector",
        proposito="Identificar piso crítico y daños severos.",
        contenido=[
            "Pisos inspeccionados, Piso crítico (número).",
            "Acceso a elementos: Todos / Casi todos / Pocos / Ninguno.",
            "Tabla: Columna, Muro concreto, Muro mampostería, Viga — Nº con daño severo/completo.",
            "Riesgo estructural severo: calculado. Si N≥1 → alerta C. Alto.",
        ],
        botones=["← Anterior", "Guardar borrador", "Subir foto", "Siguiente →", "Ir a resultado (si C)"],
        salidas=["→ INS-W4 o INS-W6"],
    )

    _screen_block(
        doc,
        titulo="INS-W4 · Paso 4 — Daño moderado piso crítico",
        ruta="/inspecciones/<id>/paso-4/",
        roles="Inspector",
        proposito="Porcentajes de daño moderado por tipo de elemento.",
        contenido=[
            "Por cada tipo: Nº examinados, Sin daño/menor, Moderado.",
            "% calculado automáticamente.",
            "Riesgo por daño moderado: A (<10%), B (10-30%), C (>30%).",
        ],
        botones=["← Anterior", "Guardar borrador", "Siguiente →"],
        salidas=["→ INS-W5"],
    )

    _screen_block(
        doc,
        titulo="INS-W5 · Paso 5 — Componentes no estructurales",
        ruta="/inspecciones/<id>/paso-5/",
        roles="Inspector",
        proposito="Losas, paredes, tanques, instalaciones, ascensores.",
        contenido=[
            "5 componentes × calificación Bajo/Medio/Alto.",
            "Riesgo de componentes: A / B / C (calculado).",
        ],
        botones=["← Anterior", "Guardar borrador", "Siguiente →"],
        salidas=["→ INS-W6"],
    )

    _screen_block(
        doc,
        titulo="INS-W6 · Paso 6 — Riesgo final y etiqueta",
        ruta="/inspecciones/<id>/paso-6/",
        roles="Inspector",
        proposito="Resultado oficial de acceso.",
        contenido=[
            "Resumen riesgos: Externo, Estructural severo, Estructural moderado, Componentes.",
            "Riesgo final: máximo de los anteriores (automático).",
            "Etiqueta: VERDE / AMARILLA / ROJA con ícono y texto normativo.",
            "Texto boletín ANIH sobre significado de cada etiqueta.",
        ],
        botones=["← Anterior", "Guardar borrador", "Siguiente →"],
        salidas=["→ INS-W7"],
    )

    _screen_block(
        doc,
        titulo="INS-W7 · Paso 7 — Acciones recomendadas",
        ruta="/inspecciones/<id>/paso-7/",
        roles="Inspector",
        proposito="Checklist de medidas.",
        contenido=[
            "Inspección detallada: Estructura / Geología / Instalaciones.",
            "Medidas: Acordonar, Cerrar calles, Apuntalar, Desconectar gas/electricidad, Otra.",
            "Campo observaciones (texto largo).",
        ],
        botones=["← Anterior", "Guardar borrador", "Siguiente →"],
        salidas=["→ INS-W8"],
    )

    _screen_block(
        doc,
        titulo="INS-W8 · Anexos — Fotos, croquis, cierre",
        ruta="/inspecciones/<id>/paso-8/",
        roles="Inspector",
        proposito="Completar evidencia y cerrar inspección.",
        contenido=[
            "Subir fotos (múltiples) de elementos dañados.",
            "Subir croquis o dibujar en canvas (opcional fase 2).",
            "Firmas: nombre inspector 1 y 2 (texto; imagen firma opcional).",
            "Nota legal de la planilla V.8.",
        ],
        botones=[
            "← Anterior",
            "Guardar borrador",
            "Finalizar inspección",
            "Generar PDF",
            "Descargar etiqueta PDF",
        ],
        validaciones=["Campos obligatorios de pasos anteriores completos"],
        salidas=[
            "→ INS-02 Detalle inspección cerrada",
            "PDF planilla + PDF etiqueta",
        ],
    )

    _screen_block(
        doc,
        titulo="INS-02 · Detalle inspección (lectura / coordinador)",
        ruta="/inspecciones/<id>/",
        roles="Inspector (propia), Coordinador (todas)",
        proposito="Consulta completa y acciones post-cierre.",
        contenido=[
            "Todas las secciones en modo lectura.",
            "Mapa con pin de coordenadas.",
            "Galería de fotos.",
            "Historial de cambios (auditoría).",
        ],
        botones=[
            "Descargar PDF planilla",
            "Descargar PDF etiqueta",
            "Enviar por correo",
            "Editar (solo si borrador o coordinador)",
            "Exportar Excel",
        ],
        salidas=["—"],
    )

    _screen_block(
        doc,
        titulo="REP-01 · Reportes",
        ruta="/reportes/",
        roles="Coordinador, Consulta",
        proposito="Análisis agregado.",
        contenido=[
            "Filtros: fechas, municipio, etiqueta, inspector.",
            "Reportes: Inscripciones por sede/sesión | Inspectores habilitados | "
            "Inspecciones por etiqueta | Mapa de calor por zona.",
        ],
        botones=["Aplicar filtros", "Exportar Excel", "Exportar PDF resumen"],
        salidas=["Descargas"],
    )

    _h(doc, "10. Flujos de proceso", 1)

    _h(doc, "10.1 Flujo — Planificación masiva de sesiones", 2)
    _p(doc, "1. Coordinador da de alta sedes una vez (SED-01). → 2. «Cargar sesiones en lote» "
         "(CAL-02): UCV, 1–15 jul, Lun–Vie, 09:00 → 11 sesiones automáticas. → "
         "3. Opcional: import Excel (CAL-03). → 4. QR impreso en cada sede. → "
         "5. Un solo enlace público /inscripcion/ o QR por sede.")

    _h(doc, "10.2 Flujo — Inscripción participante", 2)
    _p(doc, "1. Participante entra por web o QR de sede (P-02 / P-02b). → "
         "2. Elige sede (si no viene por QR) y sesión con cupo. → "
         "3. Completa datos personales (P-02c). → 4. Confirmación P-03 (+ correo opcional). → "
         "5. Día de la sesión: coordinador marca asistencia (SES-03 o SES-04). → "
         "6. Certifica. → 7. Opcional: crea usuario inspector.")

    _h(doc, "10.3 Flujo — Alta de inspector", 2)
    _p(doc, "1. Participante con certificado=True. → 2. Coordinador «Crear usuario inspector». → "
         "3. Sistema genera username (ej. cedula) y clave temporal. → 4. Correo invitación con enlace "
         "cambio de contraseña. → 5. Inspector inicia sesión L-01. → 6. Accede a INS-01.")

    _h(doc, "10.4 Flujo — Inspección en campo (wizard)", 2)
    _p(doc, "1. Inspector «Nueva inspección». → 2. Pasos 0-1 (encabezado y edificación). → "
         "3. Paso 2 externa: si Riesgo C → salto a paso 6 (etiqueta roja). → "
         "4. Si continúa: pasos 3-5. En paso 3 si N≥1 severo → salto a 6. → "
         "5. Paso 6: cálculo riesgo final y etiqueta. → 6. Paso 7 acciones. → "
         "7. Paso 8 fotos y cierre. → 8. Generar PDF planilla + etiqueta.")

    _h(doc, "10.5 Flujo — Generación PDF", 2)
    _p(doc, "1. Inspección cerrada. → 2. Motor WeasyPrint renderiza plantilla HTML planilla_erd_v8.html. → "
         "3. Inserta logos, datos, tablas, fotos, etiqueta. → 4. Guarda en media/ y ofrece descarga. → "
         "5. Opcional: envío SMTP al inspector y coordinación.")

    _h(doc, "10.6 Diagrama de decisión — Reglas de corte (paso 2 y 3)", 2)
    _table(doc, ["Condición", "Acción del sistema"], [
        ["Riesgo externo = C (paso 2)", "Ocultar pasos 3-5; ir a paso 6; sugerir etiqueta ROJA"],
        ["N elementos severo ≥ 1 (paso 3)", "Ocultar pasos 4-5; ir a paso 6; sugerir etiqueta ROJA"],
        ["Ninguna condición de corte", "Flujo completo pasos 4 y 5"],
        ["Paso 6", "riesgo_final = max(riesgo_2, riesgo_3, riesgo_4, riesgo_5)"],
        ["Etiqueta", "A→Verde, B→Amarilla, C→Roja"],
    ])

    _h(doc, "11. Correos electrónicos", 1)
    _table(doc, ["Evento", "Destinatario", "Contenido"], [
        ["Inscripción confirmada", "Participante", "Sede, fecha, hora de la sesión"],
        ["Cuenta inspector creada", "Inspector", "Usuario, enlace activación, cambio de clave"],
        ["Reset contraseña", "Usuario", "Enlace temporal Django"],
        ["Inspección cerrada", "Inspector + coordinador", "PDF adjunto (opcional)"],
        ["Recordatorio sesión", "Inscritos", "24h antes (opcional fase 2)"],
    ])
    _p(doc, "Configuración: SMTP vía variables de entorno (Brevo, Gmail app password, o servidor institucional).")

    _h(doc, "12. Seguridad y privacidad", 1)
    _bullets(doc, [
        "HTTPS obligatorio en producción.",
        "Contraseñas hasheadas (Django auth).",
        "Permisos por rol; inspector solo ve sus inspecciones.",
        "C.I., teléfono y correo: acceso restringido a roles autorizados.",
        "Registro de auditoría en inspecciones cerradas.",
        "Respaldos periódicos de PostgreSQL y carpeta media/.",
        "Consentimiento informado en inscripción pública.",
    ])

    _h(doc, "13. Despliegue y entornos", 1)
    _table(doc, ["Entorno", "BD", "Uso"], [
        ["Desarrollo (local)", "SQLite", "Programación y pruebas"],
        ["Staging", "PostgreSQL Render", "Pruebas coordinadores"],
        ["Producción", "PostgreSQL", "Operación real"],
    ])
    _p(doc, "Hosting gratuito recomendado: Render (Web Service + PostgreSQL). "
         "Alternativa: VPS Oracle Cloud Always Free con Nginx + Gunicorn.")

    _h(doc, "14. Fases de implementación", 1)
    _table(doc, ["Fase", "Duración est.", "Entregables"], [
        ["Fase 1", "2-3 semanas", "Sedes, sesiones en lote, inscripción pública (sede→sesión), QR, asistencia, roles"],
        ["Fase 2", "3-4 semanas", "Wizard inspección pasos 0-8, reglas de corte, borrador/cierre"],
        ["Fase 3", "2 semanas", "PDF planilla + etiqueta, fotos, reportes"],
        ["Fase 4", "1-2 semanas", "Correos, mapa, exportaciones, ajustes piloto"],
        ["Fase 5", "Continuo", "Mejoras, app móvil PWA, integraciones"],
    ])

    _h(doc, "15. Estructura de carpetas del proyecto", 1)
    _p(doc, "terremoto-caracas-inspecciones/")
    _bullets(doc, [
        "config/ — settings, urls, wsgi",
        "apps/capacitacion/ — campaña, sede, sesión, inscripción",
        "apps/participantes/, inspecciones/, cuentas/, reportes/",
        "templates/ — pantallas P-*, L-*, SED-*, CAL-*, SES-*, INS-W*, PDF HTML",
        "static/ — CSS Bootstrap, logos",
        "media/ — fotos y PDFs generados",
        "docs/ — este documento, diccionario datos, planilla V.8",
        "scripts/ — utilidades import/export",
        "requirements.txt",
    ])

    _h(doc, "16. Criterios de aceptación (piloto)", 1)
    _bullets(doc, [
        "Coordinador carga 20 sesiones en lote para una sede sin crear una por una.",
        "Participante inscribe eligiendo sede y sesión (o QR) con validación C.I.",
        "Día del evento: asistencia masiva desde «Sesiones de hoy» (SES-04).",
        "Inspector certificado completa inspección en móvil (responsive).",
        "Riesgo externo C salta correctamente a etiqueta roja.",
        "PDF generado coincide con estructura planilla V.8 (revisión institucional).",
        "Cada inspector solo accede a sus inspecciones.",
        "Export Excel de participantes, sesiones e inspecciones funcional.",
    ])

    _h(doc, "17. Referencias documentales", 1)
    _bullets(doc, [
        "Instrumento de Inspección — Planilla y Manual de Campo V.8.",
        "Evaluación de Daños — Boletín 61 ANIH.",
        "Manual de Entrenamiento (complemento para capacitación, no sustituye campo).",
    ])

    _h(doc, "18. Aprobaciones", 1)
    _table(doc, ["Rol", "Nombre", "Firma", "Fecha"], [
        ["Patrocinador institucional", "", "", ""],
        ["Coordinación técnica", "", "", ""],
        ["Líder desarrollo", "", "", ""],
    ])

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(OUT)
    print(f"Generado: {OUT}")


if __name__ == "__main__":
    main()
